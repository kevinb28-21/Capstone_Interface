#!/usr/bin/env python3
"""
Script to generate comprehensive technical documentation in .docx format
from all project documentation sources.
"""

import os
import re
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Project information
PROJECT_NAME = "Capstone_Interface"
PROJECT_VERSION = "1.0.0"
PROJECT_DESCRIPTION = "Drone Crop Health Dashboard - Full-stack web application for monitoring onion crop health using drone-captured images"

def setup_document():
    """Create and configure the document with professional formatting."""
    doc = Document()
    
    # Set document margins (1 inch = 914400 EMU)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Set line spacing (1.15)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph_format.line_spacing = 1.15
    
    return doc

def add_cover_page(doc):
    """Add cover page with project information."""
    # Title
    title = doc.add_heading(PROJECT_NAME, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph(PROJECT_DESCRIPTION)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0].font
    subtitle_format.size = Pt(14)
    subtitle_format.italic = True
    
    doc.add_paragraph()  # Spacing
    
    # Version
    version_para = doc.add_paragraph(f"Version {PROJECT_VERSION}")
    version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Date
    date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Page break
    doc.add_page_break()

def add_table_of_contents(doc):
    """Add table of contents with Word TOC field for automatic hyperlinks."""
    heading = doc.add_heading('Table of Contents', 1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacing
    
    # Add Word TOC field (will be updated automatically when opened in Word)
    para = doc.add_paragraph()
    run = para.add_run()
    
    # Create TOC field
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)
    run._element.append(fldChar3)
    
    doc.add_paragraph()  # Spacing after TOC
    
    # Also add manual TOC for reference (in case TOC field doesn't work)
    doc.add_paragraph("(If the table of contents above is not hyperlinked, right-click and select 'Update Field' in Microsoft Word)", style='Intense Quote')
    
    doc.add_page_break()

def read_markdown_file(file_path):
    """Read and parse markdown file, converting to plain text with formatting hints."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        return content
    except Exception as e:
        return f"[Error reading file: {str(e)}]"

def add_section_from_markdown(doc, content, heading_level=1):
    """Add content from markdown to document, preserving structure."""
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    code_language = None
    
    while i < len(lines):
        line = lines[i]
        original_line = line
        line = line.rstrip()
        
        # Check for code block start/end
        if line.strip().startswith('```'):
            if in_code_block:
                # End of code block
                if code_lines:
                    code_text = '\n'.join(code_lines)
                    para = doc.add_paragraph(code_text, style='No Spacing')
                    para_format = para.paragraph_format
                    para_format.left_indent = Inches(0.5)
                    para_format.right_indent = Inches(0.5)
                    para_format.space_before = Pt(6)
                    para_format.space_after = Pt(6)
                    para_format.keep_together = True
                    
                    # Style code block with background
                    for run in para.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0, 0, 0)
                    
                    # Add shading to code block
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), 'F5F5F5')
                    para._element.get_or_add_pPr().append(shading_elm)
                
                code_lines = []
                code_language = None
                in_code_block = False
            else:
                # Start of code block
                in_code_block = True
                code_language = line.strip()[3:].strip() if len(line.strip()) > 3 else None
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        if not line.strip():
            i += 1
            continue
        
        # Check for headings
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            heading_text = line.lstrip('#').strip()
            # Remove markdown formatting from heading
            heading_text = re.sub(r'\*\*(.*?)\*\*', r'\1', heading_text)
            heading_text = re.sub(r'\*(.*?)\*', r'\1', heading_text)
            doc.add_heading(heading_text, level=min(level, 9))
        
        # Check for lists
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            doc.add_paragraph(text, style='List Bullet')
        
        # Check for numbered lists
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            doc.add_paragraph(text, style='List Number')
        
        # Check for horizontal rules
        elif line.strip() == '---' or line.strip() == '***':
            # Skip horizontal rules in docx
            pass
        
        # Regular paragraph
        else:
            # Clean up markdown formatting
            text = line
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Bold
            text = re.sub(r'\*(.*?)\*', r'\1', text)  # Italic
            text = re.sub(r'`(.*?)`', r'\1', text)  # Inline code
            text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)  # Links
            if text.strip():
                doc.add_paragraph(text)
        
        i += 1

def add_executive_summary(doc):
    """Add executive summary section."""
    doc.add_heading('1. Executive Summary / Overview', 1)
    
    summary = f"""
{PROJECT_DESCRIPTION}

The platform performs automated analysis using vegetation indices (NDVI, SAVI, GNDVI) and machine learning models, specifically optimized for onion crop monitoring. The system provides real-time drone telemetry tracking, image upload and processing, and comprehensive crop health analytics through an interactive web dashboard.

Key Features:
• Automated image processing with vegetation index calculations
• Machine learning-based crop health classification
• Real-time drone telemetry and route tracking
• Interactive map visualization with Leaflet
• Scalable cloud storage with AWS S3
• PostgreSQL database for persistent data storage
• Background worker for automated processing pipeline

Technology Stack:
• Frontend: React with Vite, React-Leaflet for mapping
• Backend: Node.js with Express.js
• Image Processing: Python with Flask, OpenCV, TensorFlow
• Database: PostgreSQL
• Storage: AWS S3
"""
    
    add_section_from_markdown(doc, summary)

def add_getting_started(doc, project_root):
    """Add getting started section."""
    doc.add_heading('2. Getting Started / Installation Guide', 1)
    
    # Read getting started documentation
    getting_started_path = project_root / "Documentation" / "getting-started" / "README.md"
    if getting_started_path.exists():
        content = read_markdown_file(getting_started_path)
        add_section_from_markdown(doc, content, heading_level=2)
    
    # Add installation steps
    doc.add_heading('2.1 Prerequisites', 2)
    doc.add_paragraph("""
Before you begin, ensure you have the following installed:
• Node.js (v18 or higher)
• Python (v3.8 or higher)
• PostgreSQL (v12 or higher)
• Git
""")
    
    doc.add_heading('2.2 Installation Steps', 2)
    
    doc.add_heading('2.2.1 Backend Setup (Node.js)', 3)
    code_para = doc.add_paragraph("""cd server
npm install
cp .env.example .env  # Edit .env with your configuration
npm run dev""", style='No Spacing')
    code_para.paragraph_format.left_indent = Inches(0.5)
    code_para.paragraph_format.right_indent = Inches(0.5)
    code_para.paragraph_format.space_before = Pt(6)
    code_para.paragraph_format.space_after = Pt(6)
    for run in code_para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'F5F5F5')
    code_para._element.get_or_add_pPr().append(shading_elm)
    
    doc.add_heading('2.2.2 Frontend Setup (React)', 3)
    code_para = doc.add_paragraph("""cd client
npm install
npm run dev""", style='No Spacing')
    code_para.paragraph_format.left_indent = Inches(0.5)
    code_para.paragraph_format.right_indent = Inches(0.5)
    code_para.paragraph_format.space_before = Pt(6)
    code_para.paragraph_format.space_after = Pt(6)
    for run in code_para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'F5F5F5')
    code_para._element.get_or_add_pPr().append(shading_elm)
    
    doc.add_heading('2.2.3 Python Processing Service', 3)
    code_para = doc.add_paragraph("""cd python_processing
pip install -r requirements.txt
cp .env.example .env  # Edit .env with your configuration
python flask_api_db.py""", style='No Spacing')
    code_para.paragraph_format.left_indent = Inches(0.5)
    code_para.paragraph_format.right_indent = Inches(0.5)
    code_para.paragraph_format.space_before = Pt(6)
    code_para.paragraph_format.space_after = Pt(6)
    for run in code_para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'F5F5F5')
    code_para._element.get_or_add_pPr().append(shading_elm)
    
    doc.add_heading('2.2.4 Database Setup', 3)
    code_para = doc.add_paragraph("""# Create database
createdb drone_analytics

# Run schema
psql -U postgres -d drone_analytics -f server/database/schema.sql

# Run migrations (if any)
psql -U postgres -d drone_analytics -f python_processing/database_migration_add_gndvi.sql""", style='No Spacing')
    code_para.paragraph_format.left_indent = Inches(0.5)
    code_para.paragraph_format.right_indent = Inches(0.5)
    code_para.paragraph_format.space_before = Pt(6)
    code_para.paragraph_format.space_after = Pt(6)
    for run in code_para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'F5F5F5')
    code_para._element.get_or_add_pPr().append(shading_elm)

def add_architecture_section(doc, project_root):
    """Add architecture and system design section."""
    doc.add_heading('3. Architecture + System Design', 1)
    
    # Read architecture documentation
    arch_path = project_root / "Documentation" / "architecture" / "README.md"
    if arch_path.exists():
        content = read_markdown_file(arch_path)
        add_section_from_markdown(doc, content, heading_level=2)
    
    # Read project overview
    overview_path = project_root / "Documentation" / "getting-started" / "PROJECT_OVERVIEW.md"
    if overview_path.exists():
        content = read_markdown_file(overview_path)
        add_section_from_markdown(doc, content, heading_level=2)

def add_api_documentation(doc, project_root):
    """Add API documentation section."""
    doc.add_heading('4. API Documentation', 1)
    
    # Node.js API
    doc.add_heading('4.1 Node.js API', 2)
    node_api_path = project_root / "Documentation" / "api" / "NODE_API.md"
    if node_api_path.exists():
        content = read_markdown_file(node_api_path)
        add_section_from_markdown(doc, content, heading_level=3)
    
    # Flask API
    doc.add_heading('4.2 Python Flask API', 2)
    flask_api_path = project_root / "Documentation" / "api" / "FLASK_API.md"
    if flask_api_path.exists():
        content = read_markdown_file(flask_api_path)
        add_section_from_markdown(doc, content, heading_level=3)

def add_configuration_guide(doc, project_root):
    """Add configuration guide section."""
    doc.add_heading('5. Configuration Guide', 1)
    
    doc.add_heading('5.1 Environment Variables', 2)
    
    doc.add_heading('5.1.1 Backend (.env in server/)', 3)
    code_para = doc.add_paragraph("""PORT=5000
NODE_ENV=development
ORIGIN=http://localhost:5173,http://localhost:5182
AWS_ACCESS_KEY_ID=your-key
AWS_SECRET_ACCESS_KEY=your-secret
AWS_REGION=us-east-1
S3_BUCKET_NAME=your-bucket
DB_HOST=localhost
DB_PORT=5432
DB_NAME=drone_analytics
DB_USER=postgres
DB_PASSWORD=your-password""", style='No Spacing')
    code_para.paragraph_format.left_indent = Inches(0.5)
    code_para.paragraph_format.right_indent = Inches(0.5)
    code_para.paragraph_format.space_before = Pt(6)
    code_para.paragraph_format.space_after = Pt(6)
    for run in code_para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'F5F5F5')
    code_para._element.get_or_add_pPr().append(shading_elm)
    
    doc.add_heading('5.1.2 Python Processing (.env in python_processing/)', 3)
    code_para = doc.add_paragraph("""FLASK_PORT=5001
FLASK_DEBUG=True
DB_HOST=localhost
DB_PORT=5432
DB_NAME=drone_analytics
DB_USER=postgres
DB_PASSWORD=your-password
UPLOAD_FOLDER=./uploads
PROCESSED_FOLDER=./processed
WORKER_POLL_INTERVAL=10
WORKER_BATCH_SIZE=5
AWS_ACCESS_KEY_ID=your-key
AWS_SECRET_ACCESS_KEY=your-secret
AWS_REGION=us-east-1
S3_BUCKET_NAME=your-bucket""", style='No Spacing')
    code_para.paragraph_format.left_indent = Inches(0.5)
    code_para.paragraph_format.right_indent = Inches(0.5)
    code_para.paragraph_format.space_before = Pt(6)
    code_para.paragraph_format.space_after = Pt(6)
    for run in code_para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'F5F5F5')
    code_para._element.get_or_add_pPr().append(shading_elm)
    
    doc.add_heading('5.1.3 Frontend (.env in client/)', 3)
    code_para = doc.add_paragraph("""VITE_API_URL=http://localhost:5000""", style='No Spacing')
    code_para.paragraph_format.left_indent = Inches(0.5)
    code_para.paragraph_format.right_indent = Inches(0.5)
    code_para.paragraph_format.space_before = Pt(6)
    code_para.paragraph_format.space_after = Pt(6)
    for run in code_para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'F5F5F5')
    code_para._element.get_or_add_pPr().append(shading_elm)
    
    # Read deployment documentation for more config details
    deployment_path = project_root / "Documentation" / "deployment" / "README.md"
    if deployment_path.exists():
        doc.add_heading('5.2 Deployment Configuration', 2)
        content = read_markdown_file(deployment_path)
        add_section_from_markdown(doc, content, heading_level=3)

def add_usage_examples(doc, project_root):
    """Add usage examples section."""
    doc.add_heading('6. Usage Examples', 1)
    
    doc.add_heading('6.1 Image Upload Example', 2)
    code_para = doc.add_paragraph("""# Upload image via curl
curl -X POST http://localhost:5001/api/upload \\
  -F "image=@image.jpg" \\
  -F 'gps={"latitude":43.6532,"longitude":-79.3832,"altitude":100,"bearing":45,"speed":5.0}'""", style='No Spacing')
    code_para.paragraph_format.left_indent = Inches(0.5)
    code_para.paragraph_format.right_indent = Inches(0.5)
    code_para.paragraph_format.space_before = Pt(6)
    code_para.paragraph_format.space_after = Pt(6)
    for run in code_para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'F5F5F5')
    code_para._element.get_or_add_pPr().append(shading_elm)
    
    doc.add_heading('6.2 Get Processed Images', 2)
    code_para = doc.add_paragraph("""# Get all images
curl http://localhost:5001/api/data

# Get specific image
curl http://localhost:5001/api/data?image_id=uuid-here""", style='No Spacing')
    code_para.paragraph_format.left_indent = Inches(0.5)
    code_para.paragraph_format.right_indent = Inches(0.5)
    code_para.paragraph_format.space_before = Pt(6)
    code_para.paragraph_format.space_after = Pt(6)
    for run in code_para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'F5F5F5')
    code_para._element.get_or_add_pPr().append(shading_elm)
    
    doc.add_heading('6.3 Telemetry Update', 2)
    code_para = doc.add_paragraph("""curl -X POST http://localhost:5000/api/telemetry \\
  -H "Content-Type: application/json" \\
  -d '{
    "position": {"lat": 43.6532, "lng": -79.3832},
    "route": [
      {"lat": 43.6532, "lng": -79.3832},
      {"lat": 43.6542, "lng": -79.3842}
    ]
  }'""", style='No Spacing')
    code_para.paragraph_format.left_indent = Inches(0.5)
    code_para.paragraph_format.right_indent = Inches(0.5)
    code_para.paragraph_format.space_before = Pt(6)
    code_para.paragraph_format.space_after = Pt(6)
    for run in code_para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'F5F5F5')
    code_para._element.get_or_add_pPr().append(shading_elm)
    
    # Read ML training documentation
    ml_training_path = project_root / "Documentation" / "ml" / "ML_TRAINING.md"
    if ml_training_path.exists():
        doc.add_heading('6.4 ML Model Training', 2)
        content = read_markdown_file(ml_training_path)
        add_section_from_markdown(doc, content, heading_level=3)

def add_troubleshooting(doc, project_root):
    """Add troubleshooting guide section."""
    doc.add_heading('7. Troubleshooting Guide', 1)
    
    doc.add_heading('7.1 Common Issues', 2)
    
    doc.add_heading('7.1.1 Port Conflicts', 3)
    doc.add_paragraph("""
If a port is already in use:
• Backend: Set PORT in server/.env
• Frontend: Modify vite.config.js or use npm run dev -- --port <port>
• Flask: Set FLASK_PORT in python_processing/.env

Find process using port:
lsof -i :5000
kill -9 <PID>
""")
    
    doc.add_heading('7.1.2 Database Connection Errors', 3)
    doc.add_paragraph("""
• Verify PostgreSQL is running: pg_isready
• Check credentials in .env files
• Verify database exists: psql -l
• Check connection string format
""")
    
    doc.add_heading('7.1.3 S3 Upload Failures', 3)
    doc.add_paragraph("""
• Verify AWS credentials are correct
• Check bucket permissions
• Verify bucket name matches configuration
• Check AWS region setting
""")
    
    doc.add_heading('7.1.4 CORS Errors', 3)
    doc.add_paragraph("""
• Update ORIGIN in backend .env file
• Ensure frontend URL is included in allowed origins
• Check CORS configuration in server.js
""")
    
    # Read deployment troubleshooting
    troubleshooting_path = project_root / "Documentation" / "deployment" / "NETLIFY_TROUBLESHOOTING.md"
    if troubleshooting_path.exists():
        content = read_markdown_file(troubleshooting_path)
        add_section_from_markdown(doc, content, heading_level=2)

def add_contributing_guidelines(doc, project_root):
    """Add contributing guidelines section."""
    doc.add_heading('8. Contributing Guidelines', 1)
    
    # Read development documentation
    dev_path = project_root / "Documentation" / "development" / "README.md"
    if dev_path.exists():
        content = read_markdown_file(dev_path)
        add_section_from_markdown(doc, content, heading_level=2)
    
    doc.add_heading('8.1 Code Style', 2)
    doc.add_paragraph("""
• JavaScript/TypeScript: Follow ESLint configuration
• Python: Follow PEP 8 style guide
• SQL: Use uppercase for keywords
• Comments: Document complex functions and algorithms
""")
    
    doc.add_heading('8.2 Development Workflow', 2)
    doc.add_paragraph("""
1. Create a feature branch: git checkout -b feature/your-feature-name
2. Make your changes following code style guidelines
3. Test your changes locally with all services running
4. Commit and push: git commit -m "Description" && git push origin feature/your-feature-name
5. Create pull request for review
""")

def add_changelog(doc):
    """Add changelog section."""
    doc.add_heading('9. Changelog / Version History', 1)
    
    doc.add_heading('9.1 Version 1.0.0', 2)
    doc.add_paragraph("""
Initial release with the following features:
• Core image processing (NDVI, SAVI, GNDVI)
• Database integration with PostgreSQL
• S3 storage integration
• Background worker for automated processing
• Frontend dashboard with React and Leaflet
• ML model training pipeline
• Real-time telemetry tracking
• Comprehensive API documentation
""")
    
    doc.add_paragraph("""
Note: For detailed version history, refer to git commit logs and release notes.
""")

def add_appendices(doc, project_root):
    """Add appendices section."""
    doc.add_heading('10. Appendices', 1)
    
    doc.add_heading('10.1 Database Schema', 2)
    schema_path = project_root / "Documentation" / "database" / "SCHEMA.md"
    if schema_path.exists():
        content = read_markdown_file(schema_path)
        add_section_from_markdown(doc, content, heading_level=3)
    
    doc.add_heading('10.2 Background Worker Details', 2)
    bg_worker_path = project_root / "Documentation" / "ml" / "BACKGROUND_WORKER.md"
    if bg_worker_path.exists():
        content = read_markdown_file(bg_worker_path)
        add_section_from_markdown(doc, content, heading_level=3)
    
    doc.add_heading('10.3 Dataset Information', 2)
    datasets_path = project_root / "python_processing" / "datasets" / "README.md"
    if datasets_path.exists():
        content = read_markdown_file(datasets_path)
        add_section_from_markdown(doc, content, heading_level=3)
    
    doc.add_heading('10.4 S3 Integration Details', 2)
    s3_path = project_root / "Documentation" / "architecture" / "S3_INTEGRATION.md"
    if s3_path.exists():
        content = read_markdown_file(s3_path)
        add_section_from_markdown(doc, content, heading_level=3)

def main():
    """Main function to generate documentation."""
    project_root = Path(__file__).parent
    
    print("Generating technical documentation...")
    print(f"Project root: {project_root}")
    
    # Create document
    doc = setup_document()
    
    # Add sections
    print("Adding cover page...")
    add_cover_page(doc)
    
    print("Adding table of contents...")
    add_table_of_contents(doc)
    
    print("Adding executive summary...")
    add_executive_summary(doc)
    
    print("Adding getting started guide...")
    add_getting_started(doc, project_root)
    
    print("Adding architecture section...")
    add_architecture_section(doc, project_root)
    
    print("Adding API documentation...")
    add_api_documentation(doc, project_root)
    
    print("Adding configuration guide...")
    add_configuration_guide(doc, project_root)
    
    print("Adding usage examples...")
    add_usage_examples(doc, project_root)
    
    print("Adding troubleshooting guide...")
    add_troubleshooting(doc, project_root)
    
    print("Adding contributing guidelines...")
    add_contributing_guidelines(doc, project_root)
    
    print("Adding changelog...")
    add_changelog(doc)
    
    print("Adding appendices...")
    add_appendices(doc, project_root)
    
    # Save document
    output_filename = f"ProjectDocumentation_{PROJECT_NAME}.docx"
    output_path = project_root / output_filename
    
    print(f"Saving document to: {output_path}")
    doc.save(str(output_path))
    
    print(f"✓ Documentation generated successfully: {output_filename}")
    print(f"  Location: {output_path}")

if __name__ == "__main__":
    main()
